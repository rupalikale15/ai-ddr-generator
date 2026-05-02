import streamlit as st
import fitz
from groq import Groq
from docx import Document
from docx.shared import Inches
import io
import json

st.set_page_config(page_title="AI DDR Generator", layout="wide")
st.title("🛠️ AI DDR Report Generator")

# ====================== GROQ SETUP ======================
client = Groq(api_key="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")   # ← Paste your Groq API key here

# ====================== PROMPTS ======================
SYSTEM_PROMPT = """You are an expert civil engineer. Create a professional DDR by merging Inspection and Thermal data.
- Link dampness observations with thermal temperature anomalies.
- Never invent facts. Use "Not Available" if data is missing.
- Use simple, client-friendly language."""

JSON_SCHEMA = """Return ONLY valid JSON (no extra text):
{
  "property_issue_summary": "One paragraph summary",
  "area_wise_observations": [
    {
      "area": "e.g. Hall, Master Bedroom, External Wall",
      "observation": "description from inspection",
      "thermal_findings": "thermal correlation with temperatures",
      "severity": "High/Medium/Low"
    }
  ],
  "probable_root_cause": "string",
  "severity_assessment": {"level": "High/Medium/Low", "reasoning": "string"},
  "recommended_actions": ["action1", "action2"],
  "additional_notes": "string",
  "missing_information": ["item1", "item2"]
}"""

# ====================== PDF EXTRACTION ======================
def extract_text_and_images(uploaded_file):
    if uploaded_file is None:
        return {"raw_text": "", "images": []}
    
    pdf_bytes = uploaded_file.read()
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    
    full_text = ""
    images = []
    
    for page_num in range(len(doc)):
        page = doc[page_num]
        full_text += f"\n=== Page {page_num+1} ===\n{page.get_text()}\n"
        
        for img_idx, img in enumerate(page.get_images(full=True)):
            xref = img[0]
            base = doc.extract_image(xref)
            images.append({
                "page": page_num + 1,
                "image_bytes": base["image"],
                "ext": base["ext"]
            })
    return {"raw_text": full_text, "images": images}

# ====================== GENERATE DDR ======================
def generate_ddr(insp_data, therm_data):
    full_prompt = f"""
Inspection Report:
{insp_data['raw_text'][:13000]}

Thermal Report (with temperature data):
{therm_data['raw_text'][:7000]}
"""

    try:
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            temperature=0.0,
            max_tokens=2500,
            messages=[
                {"role": "system", "content": f"{SYSTEM_PROMPT}\n\n{JSON_SCHEMA}"},
                {"role": "user", "content": full_prompt}
            ]
        )
        
        content = response.choices[0].message.content.strip()
        
        if "```json" in content:
            content = content.split("```json")[1].split("```")[0].strip()
        elif "```" in content:
            content = content.split("```")[1].strip()
        
        return json.loads(content)
    except Exception as e:
        st.error(f"AI Error: {str(e)}")
        return None

# ====================== CREATE WORD REPORT (Improved Image Placement) ======================
def create_word_report(report_json, therm_images):
    doc = Document()
    doc.add_heading('Main Detailed Diagnostic Report (DDR)', 0)
    
    doc.add_heading('1. Property Issue Summary', level=1)
    doc.add_paragraph(report_json.get('property_issue_summary', 'Not Available'))
    
    doc.add_heading('2. Area-wise Observations', level=1)
    
    for i, obs in enumerate(report_json.get('area_wise_observations', [])):
        doc.add_heading(obs.get('area', f'Area {i+1}'), level=2)
        doc.add_paragraph(obs.get('observation', ''))
        doc.add_paragraph(f"Thermal Findings: {obs.get('thermal_findings', 'Not Available')}")
        doc.add_paragraph(f"Severity: {obs.get('severity', 'Not Available')}")
        
        # Improved image placement
        image_placed = False
        area_lower = obs.get('area', '').lower()
        
        for img in therm_images:
            if (i < len(therm_images)) or any(word in area_lower for word in ['hall', 'bedroom', 'kitchen', 'master', 'bathroom', 'wall', 'parking', 'external']):
                try:
                    img_stream = io.BytesIO(img["image_bytes"])
                    doc.add_picture(img_stream, width=Inches(5.5))
                    doc.add_paragraph(f"Supporting Thermal Image (Page {img['page']})")
                    image_placed = True
                    break
                except:
                    continue
                    
        if not image_placed and therm_images:
            try:
                img_stream = io.BytesIO(therm_images[0]["image_bytes"])
                doc.add_picture(img_stream, width=Inches(5.5))
                doc.add_paragraph("Supporting Thermal Image")
            except:
                doc.add_paragraph("Image Not Available")
        
        doc.add_paragraph("─" * 50)
    
    # Remaining sections
    doc.add_heading('3. Probable Root Cause', level=1)
    doc.add_paragraph(report_json.get('probable_root_cause', 'Not Available'))
    
    doc.add_heading('4. Severity Assessment', level=1)
    sev = report_json.get('severity_assessment', {})
    doc.add_paragraph(f"Level: {sev.get('level', 'Not Available')}")
    doc.add_paragraph(f"Reasoning: {sev.get('reasoning', 'Not Available')}")
    
    doc.add_heading('5. Recommended Actions', level=1)
    for action in report_json.get('recommended_actions', []):
        doc.add_paragraph(f"• {action}")
    
    doc.add_heading('6. Additional Notes', level=1)
    doc.add_paragraph(report_json.get('additional_notes', 'Not Available'))
    
    doc.add_heading('7. Missing or Unclear Information', level=1)
    for item in report_json.get('missing_information', []):
        doc.add_paragraph(f"• {item}")
    
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# ====================== UI ======================
col1, col2 = st.columns(2)
with col1:
    insp_file = st.file_uploader("Inspection Report PDF", type="pdf")
with col2:
    therm_file = st.file_uploader("Thermal Report PDF", type="pdf")

if st.button("🚀 Generate Main DDR Report", type="primary"):
    if insp_file and therm_file:
        with st.spinner("Extracting documents..."):
            insp_data = extract_text_and_images(insp_file)
            therm_data = extract_text_and_images(therm_file)
        
        with st.spinner("AI is analyzing and generating report..."):
            report_json = generate_ddr(insp_data, therm_data)
        
        if report_json:
            with st.spinner("Creating Word document with images..."):
                word_file = create_word_report(report_json, therm_data["images"])
            
            st.success("✅ Report Generated Successfully!")
            st.download_button(
                label="📥 Download Main DDR Report.docx",
                data=word_file,
                file_name="Main_DDR_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            st.json(report_json)
    else:
        st.warning("Please upload both PDF files.")

st.caption("Using Groq Llama 3.3 | Applied AI Builder Assignment")