import streamlit as st
import fitz  # PyMuPDF
from openai import OpenAI
from docx import Document
from docx.shared import Inches
import io
import json
import base64
from PIL import Image

# ====================== CONFIG ======================
st.set_page_config(page_title="AI DDR Generator", layout="wide")
st.title("🛠️ AI DDR Report Generator (Applied AI Builder Assignment)")

client = OpenAI(api_key=st.secrets.get("OPENAI_API_KEY"))

# ====================== PROMPTS ======================
SYSTEM_PROMPT = """You are a senior civil/structural engineer creating professional Detailed Diagnostic Reports.
Rules:
- Use ONLY information present in the documents. Never invent facts.
- Mention "Not Available" for missing information.
- Merge inspection observations with thermal findings logically.
- Use simple, client-friendly language."""

JSON_SCHEMA = """Return valid JSON only with this exact structure:
{
  "property_issue_summary": "One paragraph summary",
  "area_wise_observations": [
    {
      "area": "e.g. Hall, Master Bedroom, External Wall",
      "observation": "Detailed observation",
      "thermal_findings": "Thermal correlation",
      "severity": "High/Medium/Low",
      "image_note": "Which image supports this"
    }
  ],
  "probable_root_cause": "string",
  "severity_assessment": {"level": "High/Medium/Low", "reasoning": "string"},
  "recommended_actions": ["Action 1", "Action 2"],
  "additional_notes": "string",
  "missing_information": ["Item 1", "Item 2"]
}"""

# ====================== HELPERS ======================
def extract_pdf_content(pdf_file):
    doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
    full_text = ""
    images = []
    
    for page_num in range(len(doc)):
        page = doc[page_num]
        full_text += f"\n=== Page {page_num+1} ===\n{page.get_text()}\n"
        
        for img_idx, img in enumerate(page.get_images(full=True)):
            xref = img[0]
            base = doc.extract_image(xref)
            images.append({
                "page": page_num + 1,
                "bytes": base["image"],
                "ext": base["ext"],
                "index": img_idx
            })
    return full_text, images

def generate_ddr(inspection_text, thermal_text, thermal_images):
    # Caption few thermal images
    captions = []
    for img in thermal_images[:6]:  # Limit for cost/speed
        try:
            resp = client.chat.completions.create(
                model="gpt-4o",
                messages=[{
                    "role": "user",
                    "content": [
                        {"type": "text", "text": "Describe this thermal image focusing on moisture, dampness, temperature anomalies for building leakage diagnosis."},
                        {"type": "image_url", "image_url": {"url": f"data:image/{img['ext']};base64,{base64.b64encode(img['bytes']).decode()}"}}
                    ]
                }],
                max_tokens=150
            )
            captions.append(resp.choices[0].message.content)
        except:
            pass

    response = client.chat.completions.create(
        model="gpt-4o",
        temperature=0.1,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT + "\n" + JSON_SCHEMA},
            {"role": "user", "content": f"""
Inspection Report Content:
{inspection_text[:15000]}

Thermal Report Content:
{thermal_text[:8000]}

Thermal Image Descriptions:
{chr(10).join(captions)}
"""}
        ]
    )
    
    try:
        return json.loads(response.choices[0].message.content)
    except:
        st.error("Failed to parse JSON. Raw output below.")
        st.text(response.choices[0].message.content)
        return None

def create_word_report(report_json, all_images):
    doc = Document()
    doc.add_heading('Main Detailed Diagnostic Report (DDR)', 0)
    
    # 1. Summary
    doc.add_heading('1. Property Issue Summary', level=1)
    doc.add_paragraph(report_json.get('property_issue_summary', 'Not Available'))
    
    # 2. Area-wise Observations
    doc.add_heading('2. Area-wise Observations', level=1)
    for item in report_json.get('area_wise_observations', []):
        doc.add_heading(item.get('area', 'Area'), level=2)
        doc.add_paragraph(item.get('observation', ''))
        doc.add_paragraph(f"Thermal Findings: {item.get('thermal_findings', 'Not Available')}")
        doc.add_paragraph(f"Severity: {item.get('severity', 'Not Available')}")
        
        # Add first available image (you can improve matching later)
        if all_images:
            try:
                img_stream = io.BytesIO(all_images[0]["bytes"])
                doc.add_picture(img_stream, width=Inches(5.5))
                doc.add_paragraph(f"Image from page {all_images[0]['page']}")
            except:
                pass
        doc.add_paragraph("---")
    
    # 3-7. Other sections
    sections = ["probable_root_cause", "severity_assessment", "recommended_actions", "additional_notes", "missing_information"]
    titles = ["3. Probable Root Cause", "4. Severity Assessment", "5. Recommended Actions", "6. Additional Notes", "7. Missing or Unclear Information"]
    
    for key, title in zip(sections, titles):
        doc.add_heading(title, level=1)
        content = report_json.get(key, 'Not Available')
        if isinstance(content, dict):
            for k, v in content.items():
                doc.add_paragraph(f"{k}: {v}")
        elif isinstance(content, list):
            for i in content:
                doc.add_paragraph(f"• {i}")
        else:
            doc.add_paragraph(str(content))
    
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# ====================== UI ======================
col1, col2 = st.columns(2)
with col1:
    inspection_file = st.file_uploader("Inspection Report (Sample Report.pdf)", type="pdf")
with col2:
    thermal_file = st.file_uploader("Thermal Images Document", type="pdf")

if st.button("Generate Full DDR Report", type="primary", use_container_width=True):
    if inspection_file and thermal_file:
        with st.spinner("Extracting text and images..."):
            insp_text, insp_images = extract_pdf_content(inspection_file)
            therm_text, therm_images = extract_pdf_content(thermal_file)
        
        with st.spinner("GPT-4o Analyzing + Generating structured report..."):
            report_json = generate_ddr(insp_text, therm_text, therm_images)
        
        if report_json:
            with st.spinner("Creating Word document with images..."):
                word_file = create_word_report(report_json, therm_images + insp_images)
            
            st.success("✅ Report Generated Successfully!")
            st.download_button(
                label="📥 Download DDR Report.docx",
                data=word_file,
                file_name="Main_DDR_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            st.subheader("Preview (Summary)")
            st.json(report_json, expanded=False)
    else:
        st.warning("Please upload both PDFs")

st.caption("Built for Applied AI Builder Assignment | Uses GPT-4o + PyMuPDF")